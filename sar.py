# You will need to install this library: pip install python-docx
from docx import Document
from docx.shared import Pt, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH

def create_gym_form():
    doc = Document()
    
    # 1. SETUP: Set Narrow Margins (0.5 inches) to fit everything on one page
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(0.5)
        section.bottom_margin = Inches(0.5)
        section.left_margin = Inches(0.5)
        section.right_margin = Inches(0.5)

    # 2. TITLE
    title = doc.add_heading('FITNESS ASSESSMENT RECORD', level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph() # Spacer

    # 3. HELPER FUNCTION TO FORMAT TABLES
    def format_table(table):
        table.style = 'Table Grid'
        for row in table.rows:
            for cell in row.cells:
                paragraphs = cell.paragraphs
                for paragraph in paragraphs:
                    for run in paragraph.runs:
                        run.font.size = Pt(9) # Small font to fit
    
    # --- TABLE 1: VITALS ---
    doc.add_heading('1. VITALS', level=3)
    table1 = doc.add_table(rows=9, cols=6)
    headers1 = ["METRIC", "DATE", "", "", "", ""]
    metrics1 = ["HEIGHT", "WEIGHT", "TARGET WEIGHT", "BMI", "RM", "B/P", "FAT %", "VISCERAL FAT %"]
    
    # Fill Headers
    for i, text in enumerate(headers1):
        table1.rows[0].cells[i].text = text
    # Fill Metrics
    for i, metric in enumerate(metrics1):
        table1.rows[i+1].cells[0].text = metric
    format_table(table1)
    doc.add_paragraph() 

    # --- TABLE 2: BODY COMP ---
    doc.add_heading('2. BODY COMPOSITION', level=3)
    table2 = doc.add_table(rows=5, cols=7)
    headers2 = ["ZONE", "SUBCUT", "MUSCLE", "SUBCUT", "MUSCLE", "SUBCUT", "MUSCLE"]
    zones = ["WHOLE BODY", "TRUNK", "ARMS", "LEGS"]
    
    for i, text in enumerate(headers2):
        table2.rows[0].cells[i].text = text
    for i, zone in enumerate(zones):
        table2.rows[i+1].cells[0].text = zone
    format_table(table2)
    doc.add_paragraph()

    # --- TABLE 3: TEST ---
    doc.add_heading('3. FLEXIBILITY & STABILITY', level=3)
    table3 = doc.add_table(rows=4, cols=6)
    headers3 = ["SIT & REACH", "TRUNK EXT", "SLING", "SINGLE LEG", "PROPRIO.", "PLANK"]
    
    for i, text in enumerate(headers3):
        table3.rows[0].cells[i].text = text
    format_table(table3)
    doc.add_paragraph()

    # --- TABLE 4: ENDURANCE ---
    doc.add_heading('4. MUSCULAR ENDURANCE', level=3)
    table4 = doc.add_table(rows=3, cols=6)
    headers4 = ["EXERCISE", "DATE", "DATE", "DATE", "DATE", "DATE"]
    exercises = ["PUSH UP", "CURL UP"]
    
    for i, text in enumerate(headers4):
        table4.rows[0].cells[i].text = text
    for i, ex in enumerate(exercises):
        table4.rows[i+1].cells[0].text = ex
    format_table(table4)

    # Save
    doc.save('Gym_Assessment_Form.docx')
    print("Document created successfully!")

if __name__ == "__main__":
    create_gym_form()
